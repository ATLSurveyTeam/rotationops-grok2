"""
ATL Employee Scheduler 3.1 — AM engine (first Streamlit version)

Rules encoded from ATL_Employee_Scheduler_3.1:
- Fill all 43 Client Required slots.
- HARD 3:45 OPEN posts must open with 3:45 AM staff.
- START 5:45 posts start with 5:45 AM staff.
- Leads are excluded from posts and from break relief.
- Mentors + trainees stay together later; first version keeps them off split posts.
- Survey seats fill only when surveys_needed is True.
- Break relief follows Break Relief Guide (Primary → Secondary → Tertiary).
- Relief person must have a different lunch window than the person they cover.
"""

from collections import defaultdict
import pandas as pd

from fairness import sort_pool_by_fairness, load_history, is_fair_tracked


LUNCH_BY_SHIFT = {
    "3:45 AM": ["8:00 AM", "9:00 AM"],
    "5:45 AM": ["9:00 AM", "10:00 AM"],
    "12:15 PM": ["4:00 PM", "5:00 PM"],
    "2:00 PM": ["6:00 PM", "7:00 PM"],
}


def _norm(val) -> str:
    if val is None or (isinstance(val, float) and pd.isna(val)):
        return ""
    return str(val).strip()


def _yes(val) -> bool:
    return _norm(val).upper() in {"Y", "YES", "TRUE", "1"}


def _norm_pos(name: str) -> str:
    return _norm(name).replace("–", "-").replace("—", "-").replace("  ", " ").lower()


def strip_ampm(text: str) -> str:
    s = _norm(text)
    for token in (" AM", " PM", " am", " pm"):
        if s.endswith(token.strip()) or s.upper().endswith(token.strip()):
            pass
    s = s.replace(" AM", "").replace(" PM", "").replace(" am", "").replace(" pm", "")
    return s.strip()


WEEKDAYS = ["Mon", "Tue", "Wed", "Thu", "Fri", "Sat", "Sun"]


def is_off_on_date(days_off, schedule_date) -> bool:
    if schedule_date is None:
        return False
    raw = _norm(days_off)
    if not raw:
        return False
    try:
        wd = WEEKDAYS[schedule_date.weekday()]
    except Exception:
        return False
    parts = [p.strip()[:3].title() for p in raw.replace(";", ",").split(",") if p.strip()]
    return wd in parts


def filter_available(emp: pd.DataFrame, schedule_date) -> pd.DataFrame:
    if schedule_date is None or "Days Off" not in emp.columns:
        return emp
    mask = ~emp["Days Off"].apply(lambda d: is_off_on_date(d, schedule_date))
    return emp[mask].copy()


def pos_key(area, name) -> tuple:
    return (_norm(area).lower(), _norm_pos(name))


def find_assignment(assignment: dict, cand_name: str, preferred_area: str = ""):
    target = _norm_pos(cand_name)
    if "rover" in target:
        return None
    preferred = _norm(preferred_area).lower()
    same_area = None
    any_match = None
    for (area, name), rec in assignment.items():
        if name != target:
            continue
        any_match = rec
        if preferred and preferred in area:
            same_area = rec
            break
    return same_area or any_match


def pick_lunch(shift: str, counts: dict) -> str:
    options = LUNCH_BY_SHIFT.get(_norm(shift), ["12:00 PM"])
    best = min(options, key=lambda t: counts.get(t, 0))
    counts[best] = counts.get(best, 0) + 1
    return best


def other_lunches(shift: str, current: str) -> list:
    options = LUNCH_BY_SHIFT.get(_norm(shift), [])
    return [t for t in options if t and t != current]


def try_move_lunch(record: dict, avoid: str, counts: dict) -> bool:
    """Move a person's lunch to another allowed window so it does not match avoid."""
    if not record:
        return False
    current = record.get("Assigned Lunch", "")
    if current and current != avoid:
        return True
    for option in other_lunches(record.get("Shift", ""), current):
        if option != avoid:
            if current:
                counts[current] = max(0, counts.get(current, 1) - 1)
            record["Assigned Lunch"] = option
            counts[option] = counts.get(option, 0) + 1
            return True
    return False


def load_workbook_tables(path: str) -> dict:
    positions = pd.read_excel(path, sheet_name="Position Library")
    positions.columns = [str(c).strip() for c in positions.columns]
    positions = positions.dropna(subset=["Position"]).copy()

    relief = pd.read_excel(path, sheet_name="Break Relief Guide")
    relief.columns = [str(c).strip() for c in relief.columns]

    employees = pd.read_excel(path, sheet_name="Employee Master")
    employees.columns = [str(c).strip() for c in employees.columns]
    employees = employees[employees["Active"].map(_yes)].copy()

    return {
        "positions": positions,
        "relief": relief,
        "employees": employees,
    }


def _needs_divest(position_name: str) -> bool:
    return "divest" in position_name.lower()


def _needs_info(position_name: str) -> bool:
    return "information desk" in position_name.lower()


def _needs_main_inside(position_name: str) -> bool:
    return "inside center" in position_name.lower()


def _eligible_for_position(emp: pd.Series, position_name: str) -> bool:
    if _yes(emp.get("Lead")):
        return False
    if _needs_divest(position_name) and not _yes(emp.get("Divest")):
        return False
    if _needs_info(position_name) and not _yes(emp.get("Information Desk")):
        return False
    if _needs_main_inside(position_name) and not _yes(emp.get("Main Inside")):
        return False
    return True


def build_am_schedule(
    employees: pd.DataFrame,
    positions: pd.DataFrame,
    relief_guide: pd.DataFrame,
    surveys_needed: bool = True,
    schedule_date=None,
    history: pd.DataFrame | None = None,
) -> pd.DataFrame:
    emp = employees.copy()
    pos = positions.copy()
    guide = relief_guide.copy()

    emp.columns = [str(c).strip() for c in emp.columns]
    pos.columns = [str(c).strip() for c in pos.columns]
    guide.columns = [str(c).strip() for c in guide.columns]

    # Leads never take posts; honor Days Off for the selected date
    staff = emp[~emp["Lead"].map(_yes)].copy()
    staff = filter_available(staff, schedule_date)

    am_345 = staff[staff["Shift"].map(_norm) == "3:45 AM"].copy()
    am_545 = staff[staff["Shift"].map(_norm) == "5:45 AM"].copy()

    used = set()
    lunch_counts = defaultdict(int)
    assignment = {}  # position name -> dict
    if history is None:
        history = load_history()

    def take_eligible(pool: pd.DataFrame, position_name: str):
        ordered = sort_pool_by_fairness(pool, position_name, history, schedule_date)
        for idx, row in ordered.iterrows():
            eid = row["Employee ID"]
            if eid in used:
                continue
            if not _eligible_for_position(row, position_name):
                continue
            used.add(eid)
            return row
        return None

    # Sort: HARD 3:45 first, then START 5:45, then optional, surveys last
    def start_rank(rule: str) -> int:
        r = _norm(rule).upper()
        if "HARD 3:45" in r:
            return 0
        if "START 5:45" in r:
            return 1
        if "SURVEY" in r:
            return 3
        return 2

    pos["_rank"] = pos["Start of Day Rule"].map(start_rank)
    pos = pos.sort_values(["_rank", "Display Order"])

    for _, p in pos.iterrows():
        pname = _norm(p["Position"])
        area = _norm(p.get("Area"))
        rule = _norm(p.get("Start of Day Rule"))
        tier = _norm(p.get("Legacy Tier"))
        display = p.get("Display Order")

        is_survey = "survey" in area.lower() or "surveyor" in pname.lower()
        if is_survey and not surveys_needed:
            assignment[pos_key(area, pname)] = {
                "Pos #": display,
                "Area": area,
                "Position Name": pname,
                "Tier": tier or "Survey",
                "Start Rule": rule,
                "Assigned Employee": "NOT NEEDED",
                "Shift": "",
                "Assigned Lunch": "",
                "Status": "SURVEYS OFF",
                "Notes": "Survey flag is N",
            }
            continue

        person = None
        shift = ""
        notes = ""

        if "HARD 3:45" in rule.upper():
            person = take_eligible(am_345, pname)
            shift = "3:45 AM"
            notes = "HARD 3:45 OPEN"
            if person is None:
                person = take_eligible(am_545, pname)
                shift = "5:45 AM"
                notes = "HARD 3:45 OPEN — fallback 5:45"
        elif "START 5:45" in rule.upper():
            person = take_eligible(am_545, pname)
            shift = "5:45 AM"
            notes = "START 5:45"
            if person is None:
                person = take_eligible(am_345, pname)
                shift = "3:45 AM"
                notes = "START 5:45 — fallback 3:45 still on duty"
        elif is_survey:
            # Prefer survey-qualified 5:45 then 3:45
            survey_pool = pd.concat([am_545, am_345], ignore_index=True)
            survey_pool = survey_pool[survey_pool["Survey"].map(_yes)]
            person = take_eligible(survey_pool, pname)
            if person is not None:
                shift = _norm(person["Shift"])
            notes = "Survey seat"
        else:
            person = take_eligible(am_545, pname)
            if person is None:
                person = take_eligible(am_345, pname)
            if person is not None:
                shift = _norm(person["Shift"])
            notes = "Optional / zone fill"

        if person is None:
            assignment[pos_key(area, pname)] = {
                "Pos #": display,
                "Area": area,
                "Position Name": pname,
                "Tier": tier,
                "Start Rule": rule,
                "Assigned Employee": "UNFILLED",
                "Shift": "",
                "Assigned Lunch": "",
                "Status": "UNFILLED",
                "Notes": notes or "No eligible AM staff left",
            }
            continue

        lunch = pick_lunch(shift, lunch_counts)
        assignment[pos_key(area, pname)] = {
            "Pos #": display,
            "Area": area,
            "Position Name": pname,
            "Tier": tier,
            "Start Rule": rule,
            "Assigned Employee": _norm(person["Employee Name"]),
            "Employee ID": _norm(person["Employee ID"]),
            "Shift": shift,
            "Assigned Lunch": lunch,
            "Status": "ASSIGNED",
            "Notes": notes,
        }

    staff_for_rovers = staff.copy()
    am_rovers = assign_rovers(staff_for_rovers, used, ["3:45 AM", "5:45 AM"], lunch_counts, "AM")
    for r in am_rovers:
        assignment[pos_key("Rover / Relief Pool", r["Position Name"])] = r

    # Relief lookup by position name
    relief_map = {}
    for _, row in guide.iterrows():
        relief_map[_norm(row.get("Position Needing Relief"))] = [
            _norm(row.get("Primary Relief")),
            _norm(row.get("Secondary Relief")),
            _norm(row.get("Tertiary Relief")),
        ]

    rows = []
    for _, p in pos.sort_values("Display Order").iterrows():
        pname = _norm(p["Position"])
        area = _norm(p.get("Area"))
        rec = dict(assignment.get(pos_key(area, pname), {}))
        if not rec:
            continue

        candidates = [c for c in relief_map.get(pname, []) if c]
        relief_name = ""
        relief_from = ""
        relief_lunch = ""
        status = rec.get("Status", "")

        if rec.get("Assigned Employee") in {"UNFILLED", "NOT NEEDED"}:
            rec["Break Relief Agent"] = ""
            rec["Relief From Position"] = ""
            rec["Relief Lunch"] = ""
            rows.append(rec)
            continue

        if not candidates or candidates[0].lower().startswith("no relief"):
            rec["Break Relief Agent"] = "No Relief Needed"
            rec["Relief From Position"] = ""
            rec["Relief Lunch"] = ""
            rec["Status"] = "OK"
            rows.append(rec)
            continue

        primary_lunch = rec.get("Assigned Lunch", "")
        found = False
        for cand_pos in candidates:
            if "rover" in cand_pos.lower():
                rover_list = [v for v in assignment.values() if str(v.get("Position Name","")).startswith("Rover")]
                rover = _pick_named_rover(rover_list, primary_lunch)
                if rover is not None:
                    rec["Break Relief Agent"] = rover.get("Assigned Employee", "")
                    rec["Relief From Position"] = rover.get("Position Name", "Rover")
                    rec["Relief Lunch"] = rover.get("Assigned Lunch", "")
                    rec["Status"] = "OK — Rover"
                    found = True
                    break
            other = find_assignment(assignment, cand_pos, area)
            if not other:
                continue
            other_name = other.get("Assigned Employee", "")
            if other_name in {"", "UNFILLED", "NOT NEEDED"}:
                continue
            other_lunch = other.get("Assigned Lunch", "")
            if other_lunch and other_lunch == primary_lunch:
                if try_move_lunch(other, primary_lunch, lunch_counts):
                    other_lunch = other.get("Assigned Lunch", "")
                elif try_move_lunch(rec, other_lunch, lunch_counts):
                    primary_lunch = rec.get("Assigned Lunch", "")
                else:
                    continue
            relief_name = other_name
            relief_from = cand_pos
            relief_lunch = other_lunch
            found = True
            break

        if not rec.get("Break Relief Agent"):
            rec["Break Relief Agent"] = relief_name
            rec["Relief From Position"] = relief_from
            rec["Relief Lunch"] = relief_lunch
        if rec.get("Status") not in {"OK", "OK — Rover"}:
            if found and rec.get("Break Relief Agent"):
                rec["Status"] = "OK"
            elif not found:
                rec["Status"] = "NO RELIEF / SAME LUNCH"
                rec["Notes"] = (rec.get("Notes", "") + " | Relief chain blocked").strip(" |")

        rows.append(rec)

    for r in am_rovers:
        rows.append(r)

    cols = [
        "Pos #", "Area", "Position Name", "Tier", "Start Rule",
        "Assigned Employee", "Shift", "Assigned Lunch",
        "Break Relief Agent", "Relief From Position", "Relief Lunch",
        "Status", "Notes",
    ]
    out = pd.DataFrame(rows)
    for c in cols:
        if c not in out.columns:
            out[c] = ""
    return out[cols]


def build_schedule(employees: pd.DataFrame, positions: pd.DataFrame, relief_guide: pd.DataFrame | None = None, surveys_needed: bool = True) -> pd.DataFrame:
    """Back-compatible wrapper used by app.py."""
    if relief_guide is None:
        relief_guide = pd.DataFrame(columns=["Position Needing Relief", "Primary Relief", "Secondary Relief", "Tertiary Relief"])
    return build_am_schedule(employees, positions, relief_guide, surveys_needed=surveys_needed)

def _handoff_shift(pm_handoff: str, am_shift: str) -> str:
    text = _norm(pm_handoff).lower()
    if "not staffed" in text:
        return "NOT STAFFED"
    if "survey" in text:
        return "SURVEY"
    if "12:15" in text:
        return "12:15 PM"
    if "2:00" in text:
        return "2:00 PM"
    # Covered by either — follow AM person still on post
    if "3:45" in _norm(am_shift):
        return "12:15 PM"
    return "2:00 PM"


def build_pm_schedule(
    employees: pd.DataFrame,
    positions: pd.DataFrame,
    relief_guide: pd.DataFrame,
    am_board: pd.DataFrame,
    surveys_needed: bool = True,
    schedule_date=None,
    history: pd.DataFrame | None = None,
) -> pd.DataFrame:
    emp = employees.copy()
    pos = positions.copy()
    staff = emp[~emp["Lead"].map(_yes)].copy()
    staff = filter_available(staff, schedule_date)

    pm_1215 = staff[staff["Shift"].map(_norm) == "12:15 PM"].copy()
    pm_200 = staff[staff["Shift"].map(_norm) == "2:00 PM"].copy()

    used = set()
    lunch_counts = defaultdict(int)
    if history is None:
        history = load_history()

    def am_lookup(area, pname):
        target = _norm_pos(pname)
        area_l = _norm(area).lower()
        for _, r in am_board.iterrows():
            if _norm_pos(r.get("Position Name","")) != target:
                continue
            if _norm(r.get("Area")).lower() == area_l:
                return r
        return {}

    def take_eligible(pool: pd.DataFrame, position_name: str):
        ordered = sort_pool_by_fairness(pool, position_name, history, schedule_date)
        for idx, row in ordered.iterrows():
            eid = row["Employee ID"]
            if eid in used:
                continue
            if not _eligible_for_position(row, position_name):
                continue
            used.add(eid)
            return row
        return None

    assignment = {}
    pos = pos.sort_values("Display Order")

    for _, p in pos.iterrows():
        pname = _norm(p["Position"])
        area = _norm(p.get("Area"))
        tier = _norm(p.get("Legacy Tier"))
        display = p.get("Display Order")
        handoff = _norm(p.get("PM Handoff"))
        am_row = am_lookup(area, pname)
        am_shift = _norm(am_row.get("Shift", ""))
        am_name = _norm(am_row.get("Assigned Employee", ""))

        is_survey = "survey" in area.lower() or "surveyor" in pname.lower()
        wanted = _handoff_shift(handoff, am_shift)

        if wanted == "NOT STAFFED":
            assignment[pos_key(area, pname)] = {
                "Pos #": display, "Area": area, "Position Name": pname, "Tier": tier,
                "PM Handoff Rule": handoff, "Assigned Employee": "NOT STAFFED",
                "Shift": "", "Assigned Lunch": "", "Status": "NOT STAFFED PM",
                "Notes": f"AM was {am_name}" if am_name else "Not staffed in PM",
            }
            continue

        if is_survey and not surveys_needed:
            assignment[pos_key(area, pname)] = {
                "Pos #": display, "Area": area, "Position Name": pname, "Tier": tier or "Survey",
                "PM Handoff Rule": handoff, "Assigned Employee": "NOT NEEDED",
                "Shift": "", "Assigned Lunch": "", "Status": "SURVEYS OFF",
                "Notes": "PM survey flag is N",
            }
            continue

        person = None
        shift = wanted if wanted in {"12:15 PM", "2:00 PM"} else "12:15 PM"

        if is_survey:
            survey_pool = pd.concat([pm_1215, pm_200], ignore_index=True)
            survey_pool = survey_pool[survey_pool["Survey"].map(_yes)]
            person = take_eligible(survey_pool, pname)
            if person is not None:
                shift = _norm(person["Shift"])
        elif shift == "12:15 PM":
            person = take_eligible(pm_1215, pname)
            if person is None:
                person = take_eligible(pm_200, pname)
                if person is not None:
                    shift = "2:00 PM"
        else:
            person = take_eligible(pm_200, pname)
            if person is None:
                person = take_eligible(pm_1215, pname)
                if person is not None:
                    shift = "12:15 PM"

        if person is None:
            assignment[pos_key(area, pname)] = {
                "Pos #": display, "Area": area, "Position Name": pname, "Tier": tier,
                "PM Handoff Rule": handoff, "Assigned Employee": "UNFILLED",
                "Shift": "", "Assigned Lunch": "", "Status": "UNFILLED",
                "Notes": f"Handoff from AM {am_name} ({am_shift})" if am_name else handoff,
            }
            continue

        lunch = pick_lunch(shift, lunch_counts)
        assignment[pos_key(area, pname)] = {
            "Pos #": display, "Area": area, "Position Name": pname, "Tier": tier,
            "PM Handoff Rule": handoff,
            "Assigned Employee": _norm(person["Employee Name"]),
            "Employee ID": _norm(person["Employee ID"]),
            "Shift": shift, "Assigned Lunch": lunch, "Status": "ASSIGNED",
            "Notes": f"Handoff from {am_name} ({am_shift})" if am_name else handoff,
        }

    staff_for_rovers = staff.copy()
    pm_rovers = assign_rovers(staff_for_rovers, used, ["12:15 PM", "2:00 PM"], lunch_counts, "PM")
    for r in pm_rovers:
        assignment[pos_key("Rover / Relief Pool", r["Position Name"])] = r

    relief_map = {}
    for _, row in relief_guide.iterrows():
        relief_map[_norm(row.get("Position Needing Relief"))] = [
            _norm(row.get("Primary Relief")),
            _norm(row.get("Secondary Relief")),
            _norm(row.get("Tertiary Relief")),
        ]

    rows = []
    for _, p in pos.iterrows():
        pname = _norm(p["Position"])
        area = _norm(p.get("Area"))
        rec = dict(assignment.get(pos_key(area, pname), {}))
        if not rec:
            continue
        if rec.get("Assigned Employee") in {"UNFILLED", "NOT NEEDED", "NOT STAFFED"}:
            rec["Break Relief Agent"] = ""
            rec["Relief From Position"] = ""
            rec["Relief Lunch"] = ""
            rows.append(rec)
            continue

        candidates = [c for c in relief_map.get(pname, []) if c]
        if not candidates or candidates[0].lower().startswith("no relief"):
            rec["Break Relief Agent"] = "No Relief Needed"
            rec["Relief From Position"] = ""
            rec["Relief Lunch"] = ""
            rec["Status"] = "OK"
            rows.append(rec)
            continue

        primary_lunch = rec.get("Assigned Lunch", "")
        found = False
        for cand_pos in candidates:
            if cand_pos.lower() == "rover/relief":
                rover = _pick_named_rover(pm_rovers, primary_lunch)
                if rover is not None:
                    rec["Break Relief Agent"] = rover["Assigned Employee"]
                    rec["Relief From Position"] = rover["Position Name"]
                    rec["Relief Lunch"] = rover.get("Assigned Lunch", "")
                    rec["Status"] = "OK — Rover"
                    found = True
                    break
            other = find_assignment(assignment, cand_pos, area)
            if not other:
                continue
            other_name = other.get("Assigned Employee", "")
            if other_name in {"", "UNFILLED", "NOT NEEDED", "NOT STAFFED"}:
                continue
            other_lunch = other.get("Assigned Lunch", "")
            if other_lunch and other_lunch == primary_lunch:
                if try_move_lunch(other, primary_lunch, lunch_counts):
                    other_lunch = other.get("Assigned Lunch", "")
                elif try_move_lunch(rec, other_lunch, lunch_counts):
                    primary_lunch = rec.get("Assigned Lunch", "")
                else:
                    continue
            rec["Break Relief Agent"] = other_name
            rec["Relief From Position"] = cand_pos
            rec["Relief Lunch"] = other_lunch
            rec["Status"] = "OK"
            found = True
            break
        if not found:
            rec["Break Relief Agent"] = ""
            rec["Relief From Position"] = ""
            rec["Relief Lunch"] = ""
            rec["Status"] = "NO RELIEF / SAME LUNCH"
        rows.append(rec)

    for r in pm_rovers:
        rows.append(r)

    cols = [
        "Pos #", "Area", "Position Name", "Tier", "PM Handoff Rule",
        "Assigned Employee", "Shift", "Assigned Lunch",
        "Break Relief Agent", "Relief From Position", "Relief Lunch",
        "Status", "Notes",
    ]
    out = pd.DataFrame(rows)
    for c in cols:
        if c not in out.columns:
            out[c] = ""
    return out[cols]


def build_full_day(employees, positions, relief_guide, am_surveys=True, pm_surveys=True, schedule_date=None, history=None):
    if history is None:
        history = load_history()
    am = build_am_schedule(employees, positions, relief_guide, surveys_needed=am_surveys, schedule_date=schedule_date, history=history)
    pm = build_pm_schedule(employees, positions, relief_guide, am, surveys_needed=pm_surveys, schedule_date=schedule_date, history=history)
    return am, pm


def _remaining_pool(staff, used, shifts):
    pool = staff[staff["Shift"].map(_norm).isin(shifts)].copy()
    pool = pool[~pool["Employee ID"].isin(used)]
    return pool


def assign_rovers(staff, used, shifts, lunch_counts, prefix):
    """Turn leftover staff into named Rover #1, #2, ..."""
    pool = _remaining_pool(staff, used, shifts)
    rovers = []
    n = 1
    for _, row in pool.iterrows():
        used.add(row["Employee ID"])
        shift = _norm(row["Shift"])
        lunch = pick_lunch(shift, lunch_counts)
        name = f"Rover (#{n})"
        rec = {
            "Pos #": 100 + n,
            "Area": "Rover / Relief Pool",
            "Position Name": name,
            "Tier": "Rover",
            "Assigned Employee": _norm(row["Employee Name"]),
            "Employee ID": _norm(row["Employee ID"]),
            "Shift": shift,
            "Assigned Lunch": lunch,
            "Status": "OK",
            "Notes": "Named rover",
            "Break Relief Agent": "No Relief Needed",
            "Relief From Position": "",
            "Relief Lunch": "",
        }
        rovers.append(rec)
        n += 1
        if n > 12:
            break
    return rovers


def _pick_named_rover(rovers, primary_lunch):
    for r in rovers:
        if r.get("Assigned Lunch") and r.get("Assigned Lunch") == primary_lunch:
            continue
        if r.get("Assigned Employee"):
            return r
    return rovers[0] if rovers else None


OFFICIAL_SECTIONS = [
    ("MAIN SECURITY CHECKPOINT", [
        "Front Entrance (#1)",
        "Front Entrance (#2)",
        "ADA Entrance (1)",
        "Lanes 17-18 (1)",
        "Inside Center (1)",
        "Lanes 1-3/Pre-Check (1)",
    ]),
    ("ATRIUM", ["KIA (1)"]),
    ("SOUTH SECURITY CHECKPOINT", [
        "Checkpoint - Front Entrance (#1)",
        "Checkpoint – Front Entrance (#1)",
        "Checkpoint - Front Entrance (#2)",
        "Checkpoint – Front Entrance (#2)",
        "Checkpoint - Inside (#1)",
        "Checkpoint – Inside (#1)",
        "Checkpoint - Inside (#2)",
        "Checkpoint – Inside (#2)",
        "Hallway Pre-check Overflow (1)",
        "South Restroom (1)",
        "IHOP (1)",
    ]),
    ("NORTH SECURITY CHECKPOINT", [
        "Checkpoint - Inside (1)",
        "Checkpoint – Inside (1)",
        "North Corridor (1)",
        "North Restroom Hallway (#1)",
        "North Restroom Hallway (#2)",
        "We Juice It (1)",
    ]),
    ("LOWER NORTH SECURITY CHECKPOINT", [
        "Checkpoint - Inside (1)",
        "Checkpoint – Inside (1)",
        "N5 Door (1)",
        "LN2 Elevator (1)",
    ]),
    ("AGT LEVEL", [
        "T Platform Train Doors (#1)",
        "T Platform Train Doors (#2)",
        "Domestic Baggage Claim (1)",
        "T Bridge Landing (#1)",
        "T Bridge Landing (#2)",
        "Alpha West (1)",
        "Bravo West (1)",
        "Charlie West (1)",
        "Delta West (1)",
    ]),
    ("SURVEYS", [
        "Surveyor (#1)", "Surveyor (#2)", "Surveyor (#3)", "Surveyor (#4)",
    ]),
    ("INFORMATION DESK", ["Information Desk Personnel"]),
    ("DIVESTING", [
        "North Divesting (#1)",
        "North Divesting (#2)",
        "North Divesting (#3)",
        "North Divesting (#4)",
        "Lower North Divesting (#1)",
        "Lower North Divesting (#2)",
        "Lower North Divesting (#3)",
    ]),
]


def _lookup_row(board, wanted_names):
    by = {_norm(r["Position Name"]).replace("–", "-").replace("—", "-"): r for _, r in board.iterrows()}
    for name in wanted_names:
        key = _norm(name).replace("–", "-").replace("—", "-")
        if key in by:
            return by[key]
    return None


def build_official_sheet(am_board, pm_board, schedule_date, leads=""):
    """Lead-familiar CXR Daily Assignment Sheet rows."""
    rows = []
    used_am = set()
    used_pm = set()

    # Unique positions in official order, avoiding hyphen-variant duplicates
    seen = set()
    official_unique = []
    for section, names in [
        ("MAIN SECURITY CHECKPOINT", [
            "Front Entrance (#1)", "Front Entrance (#2)", "ADA Entrance (1)",
            "Lanes 17-18 (1)", "Inside Center (1)", "Lanes 1-3/Pre-Check (1)",
        ]),
        ("ATRIUM", ["KIA (1)"]),
        ("SOUTH SECURITY CHECKPOINT", [
            "Checkpoint - Front Entrance (#1)", "Checkpoint - Front Entrance (#2)",
            "Checkpoint - Inside (#1)", "Checkpoint - Inside (#2)",
            "Hallway Pre-check Overflow (1)", "South Restroom (1)", "IHOP (1)",
        ]),
        ("NORTH SECURITY CHECKPOINT", [
            "Checkpoint - Inside (1)", "North Corridor (1)",
            "North Restroom Hallway (#1)", "North Restroom Hallway (#2)", "We Juice It (1)",
        ]),
        ("LOWER NORTH SECURITY CHECKPOINT", [
            "Checkpoint - Inside (1)", "N5 Door (1)", "LN2 Elevator (1)",
        ]),
        ("AGT LEVEL", [
            "T Platform Train Doors (#1)", "T Platform Train Doors (#2)",
            "Domestic Baggage Claim (1)", "T Bridge Landing (#1)", "T Bridge Landing (#2)",
            "Alpha West (1)", "Bravo West (1)", "Charlie West (1)", "Delta West (1)",
        ]),
        ("SURVEYS", ["Surveyor (#1)", "Surveyor (#2)", "Surveyor (#3)", "Surveyor (#4)"]),
        ("INFORMATION DESK", ["Information Desk Personnel"]),
        ("DIVESTING", [
            "North Divesting (#1)", "North Divesting (#2)", "North Divesting (#3)", "North Divesting (#4)",
            "Lower North Divesting (#1)", "Lower North Divesting (#2)", "Lower North Divesting (#3)",
        ]),
    ]:
        official_unique.append((section, names))

    def find_row(board, pname, used):
        target = _norm(pname).replace("–", "-").replace("—", "-").lower()
        for _, r in board.iterrows():
            key = _norm(r["Position Name"]).replace("–", "-").replace("—", "-").lower()
            if key == target:
                ident = (key, _norm(r.get("Area")))
                # Lower North vs North both have Checkpoint - Inside (1)
                return r
        return None

    # More precise: match by Area + position
    def find_row_area(board, section, pname):
        target = _norm(pname).replace("–", "-").replace("—", "-").lower()
        section_key = section.lower()
        for _, r in board.iterrows():
            key = _norm(r["Position Name"]).replace("–", "-").replace("—", "-").lower()
            area = _norm(r.get("Area")).lower()
            if key != target:
                continue
            if "lower north" in section_key and "lower north" in area:
                return r
            if "north security" in section_key and "north security" in area and "lower" not in area:
                return r
            if "south" in section_key and "south" in area:
                return r
            if "main" in section_key and "main" in area:
                return r
            if "atrium" in section_key and "atrium" in area:
                return r
            if "agt" in section_key and "agt" in area:
                return r
            if "survey" in section_key and "survey" in area:
                return r
            if "information" in section_key and "information" in area:
                return r
            if "divest" in section_key and "divest" in key:
                return r
        # fallback by name only
        for _, r in board.iterrows():
            key = _norm(r["Position Name"]).replace("–", "-").replace("—", "-").lower()
            if key == target:
                return r
        return None

    out = []
    for section, names in official_unique:
        out.append({
            "Name (AM)": "", "Shift (AM)": "", "Lunch (AM)": "",
            "POSITIONS": section,
            "Name (PM)": "", "Shift (PM)": "", "Lunch (PM)": "",
            "_header": True,
        })
        for pname in names:
            am = find_row_area(am_board, section, pname)
            pm = find_row_area(pm_board, section, pname)
            out.append({
                "Name (AM)": "" if am is None else _norm(am.get("Assigned Employee")),
                "Shift (AM)": "" if am is None else _norm(am.get("Shift")),
                "Lunch (AM)": "" if am is None else _norm(am.get("Assigned Lunch")),
                "POSITIONS": pname.replace(" - ", " – "),
                "Name (PM)": "" if pm is None else _norm(pm.get("Assigned Employee")),
                "Shift (PM)": "" if pm is None else _norm(pm.get("Shift")),
                "Lunch (PM)": "" if pm is None else _norm(pm.get("Assigned Lunch")),
                "_header": False,
            })

    # Rover section
    out.append({
        "Name (AM)": "", "Shift (AM)": "", "Lunch (AM)": "",
        "POSITIONS": "ROVER / RELIEF POOL",
        "Name (PM)": "", "Shift (PM)": "", "Lunch (PM)": "",
        "_header": True,
    })
    am_rovers = am_board[am_board["Area"] == "Rover / Relief Pool"] if "Area" in am_board.columns else pd.DataFrame()
    pm_rovers = pm_board[pm_board["Area"] == "Rover / Relief Pool"] if "Area" in pm_board.columns else pd.DataFrame()
    max_r = max(len(am_rovers), len(pm_rovers))
    for i in range(max_r):
        am = am_rovers.iloc[i] if i < len(am_rovers) else None
        pm = pm_rovers.iloc[i] if i < len(pm_rovers) else None
        out.append({
            "Name (AM)": "" if am is None else _norm(am.get("Assigned Employee")),
            "Shift (AM)": "" if am is None else _norm(am.get("Shift")),
            "Lunch (AM)": "" if am is None else _norm(am.get("Assigned Lunch")),
            "POSITIONS": f"Rover (#{i+1})",
            "Name (PM)": "" if pm is None else _norm(pm.get("Assigned Employee")),
            "Shift (PM)": "" if pm is None else _norm(pm.get("Shift")),
            "Lunch (PM)": "" if pm is None else _norm(pm.get("Assigned Lunch")),
            "_header": False,
        })

    df = pd.DataFrame(out)
    return df


def _board_index(board: pd.DataFrame):
    idx = {}
    for _, r in board.iterrows():
        key = (_norm_pos(r.get("Position Name", "")), _norm(r.get("Area")).lower())
        idx[key] = r
        idx.setdefault(_norm_pos(r.get("Position Name", "")), r)
    return idx


def _match_official_row(index, section, position):
    pos = _norm_pos(position)
    section_l = _norm(section).lower()
    # try area-aware first
    for key, row in list(index.items()):
        if not (isinstance(key, tuple) and len(key) == 2):
            continue
        pkey, area = key
        if pkey != pos:
            continue
        if "lower north" in section_l and "lower north" in area:
            return row
        if section_l.startswith("north") and "north security" in area and "lower" not in area:
            return row
        if "south" in section_l and "south" in area:
            return row
        if "main" in section_l and "main" in area:
            return row
        if "atrium" in section_l and "atrium" in area:
            return row
        if "agt" in section_l and "agt" in area:
            return row
        if "survey" in section_l and "survey" in area:
            return row
        if "information" in section_l and "information" in area:
            return row
        if "divest" in section_l and "divest" in pkey:
            return row
    return index.get(pos)


def fill_official_docx(template_path: str, am_board: pd.DataFrame, pm_board: pd.DataFrame,
                       schedule_date, leads: str = "", pax: str = "") -> bytes:
    from copy import copy
    from io import BytesIO
    from docx import Document

    doc = Document(template_path)
    if isinstance(schedule_date, str):
        date_text = schedule_date
    else:
        date_text = schedule_date.strftime("%A %B %d, %Y")

    if doc.tables:
        meta = doc.tables[0]
        if len(meta.rows) >= 1 and len(meta.columns) >= 2:
            meta.cell(0, 1).text = date_text
        if len(meta.rows) >= 2 and len(meta.columns) >= 2:
            meta.cell(1, 1).text = leads or ""
        if len(meta.rows) >= 3 and len(meta.columns) >= 2:
            meta.cell(2, 1).text = pax or ""

    am_idx = _board_index(am_board)
    pm_idx = _board_index(pm_board)
    section = ""
    skip = {"positions", "surveys", "information desk", "divesting"}

    sheet = doc.tables[1]
    for row in sheet.rows:
        pos_text = _norm(row.cells[3].text)
        key = _norm_pos(pos_text)
        if not key:
            continue
        if key in skip:
            continue
        # section headers are merged across the row
        if pos_text == _norm(row.cells[0].text) and pos_text.upper() == pos_text or pos_text.isupper():
            if pos_text not in {"POSITIONS"}:
                section = pos_text
            continue
        if key in {
            "main security checkpoint", "atrium", "south security checkpoint",
            "north security checkpoint", "lower north security checkpoint",
            "agt level", "surveys", "information desk", "divesting",
        }:
            section = pos_text
            continue

        am = _match_official_row(am_idx, section, pos_text)
        pm = _match_official_row(pm_idx, section, pos_text)

        def val(rowobj, field):
            if rowobj is None:
                return ""
            v = _norm(rowobj.get(field, ""))
            if v in {"UNFILLED", "NOT NEEDED", "NOT STAFFED", "NOT STAFFED PM"}:
                return v
            return v

        row.cells[0].text = val(am, "Assigned Employee")
        row.cells[1].text = strip_ampm(val(am, "Shift"))
        row.cells[2].text = strip_ampm(val(am, "Assigned Lunch"))
        row.cells[4].text = val(pm, "Assigned Employee")
        row.cells[5].text = strip_ampm(val(pm, "Shift"))
        row.cells[6].text = strip_ampm(val(pm, "Assigned Lunch"))

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def apply_lead_sheet_overrides(official_edit: pd.DataFrame, am_board: pd.DataFrame,
                               pm_board: pd.DataFrame, employees: pd.DataFrame) -> tuple:
    """Lock typed names from the official lead sheet, fill blanks from rovers, keep rest stable."""
    am = am_board.copy()
    pm = pm_board.copy()
    emp_by_name = {}
    for _, row in employees.iterrows():
        emp_by_name[_norm(row.get("Employee Name")).lower()] = row

    def lookup_emp(name):
        return emp_by_name.get(_norm(name).lower())

    def set_name(board, section, pname, new_name, side):
        new_name = _norm(new_name)
        if new_name.upper() in {"", "UNFILLED", "NOT NEEDED", "NOT STAFFED", "NOT STAFFED PM"}:
            target_blank = True
        else:
            target_blank = False
        target = _norm_pos(pname)
        section_l = _norm(section).lower()
        idx_hit = None
        for i, row in board.iterrows():
            if _norm_pos(row.get("Position Name", "")) != target:
                continue
            area = _norm(row.get("Area")).lower()
            if "lower north" in section_l and "lower north" not in area:
                continue
            if section_l.startswith("north") and "lower" not in section_l and "lower" in area:
                continue
            if "south" in section_l and "south" not in area:
                continue
            if "main" in section_l and "main" not in area and "front" in target:
                continue
            idx_hit = i
            break
        if idx_hit is None:
            for i, row in board.iterrows():
                if _norm_pos(row.get("Position Name", "")) == target:
                    idx_hit = i
                    break
        if idx_hit is None:
            return board
        if target_blank:
            board.at[idx_hit, "Assigned Employee"] = "UNFILLED"
            board.at[idx_hit, "Shift"] = ""
            board.at[idx_hit, "Assigned Lunch"] = ""
            board.at[idx_hit, "Status"] = "UNFILLED"
            board.at[idx_hit, "Notes"] = "Cleared by override"
            return board
        info = lookup_emp(new_name)
        board.at[idx_hit, "Assigned Employee"] = new_name
        if info is not None:
            board.at[idx_hit, "Shift"] = _norm(info.get("Shift"))
            if "Assigned Lunch" in board.columns and not _norm(board.at[idx_hit, "Assigned Lunch"]):
                board.at[idx_hit, "Assigned Lunch"] = pick_lunch(_norm(info.get("Shift")), defaultdict(int))
        board.at[idx_hit, "Status"] = "LOCKED OVERRIDE"
        board.at[idx_hit, "Notes"] = "Locked from lead sheet"
        return board

    current_section = ""
    for _, row in official_edit.iterrows():
        pos = _norm(row.get("POSITIONS"))
        if not pos:
            continue
        if pos.isupper() or pos in {"ROVER / RELIEF POOL"}:
            current_section = pos
            continue
        am_name = _norm(row.get("Name (AM)"))
        pm_name = _norm(row.get("Name (PM)"))
        am = set_name(am, current_section, pos, am_name, "AM")
        pm = set_name(pm, current_section, pos, pm_name, "PM")

    def fill_from_rovers(board):
        rover_idx = [
            i for i, r in board.iterrows()
            if _norm(r.get("Area")) == "Rover / Relief Pool" and _norm(r.get("Assigned Employee")) not in {"", "UNFILLED"}
        ]
        open_idx = [
            i for i, r in board.iterrows()
            if _norm(r.get("Assigned Employee")) in {"", "UNFILLED"}
            and _norm(r.get("Area")) != "Rover / Relief Pool"
            and "survey" not in _norm(r.get("Position Name")).lower()
        ]
        for oi in open_idx:
            if not rover_idx:
                break
            ri = rover_idx.pop(0)
            board.at[oi, "Assigned Employee"] = board.at[ri, "Assigned Employee"]
            board.at[oi, "Shift"] = board.at[ri, "Shift"]
            board.at[oi, "Assigned Lunch"] = board.at[ri, "Assigned Lunch"]
            board.at[oi, "Status"] = "FILLED FROM ROVER"
            board.at[oi, "Notes"] = f"From {board.at[ri, 'Position Name']}"
            board.at[ri, "Assigned Employee"] = "UNFILLED"
            board.at[ri, "Status"] = "ROVER USED"
        return board

    am = fill_from_rovers(am)
    pm = fill_from_rovers(pm)
    return am, pm
